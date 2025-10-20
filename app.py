"""
Smart ATS Resume Checker - Enhanced v4.0 FIXED
Complete Production-Ready Backend with Groq AI Integration

ENHANCEMENTS:
- Fixed PDF formatting (matches preview exactly)
- Added LinkedIn, Languages, Projects sections
- Enhanced grammar correction
- Better keyword integration from JD
- Improved ATS scoring
- ALL ERRORS FIXED
"""

from flask import Flask, request, jsonify, session, send_file
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import os
import sqlite3
import PyPDF2
import docx
import re
import json
from groq import Groq
from datetime import timedelta, datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
import logging

load_dotenv()

app = Flask(__name__, static_folder='.')
app.secret_key = os.getenv('SECRET_KEY', 'dev-secret-key-change-in-production-12345')
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['GENERATED_FOLDER'] = 'generated'
app.config['MAX_CONTENT_LENGTH'] = 5242880
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = False
app.config['SESSION_COOKIE_HTTPONLY'] = True

CORS(app, supports_credentials=True, resources={r"/api/*": {"origins": "*"}})
limiter = Limiter(app=app, key_func=get_remote_address, default_limits=["200 per day", "50 per hour"], storage_uri="memory://")

GROQ_API_KEY = os.getenv('GROQ_API_KEY')
if GROQ_API_KEY:
    try:
        groq_client = Groq(api_key=GROQ_API_KEY)
        print("✓ Groq AI configured")
    except Exception as e:
        print(f"⚠️  Groq initialization error: {e}")
        groq_client = None
else:
    print("⚠️  GROQ_API_KEY not found")
    groq_client = None

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['GENERATED_FOLDER'], exist_ok=True)
os.makedirs('database', exist_ok=True)

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

def get_db():
    db = sqlite3.connect('database/ats_checker.db')
    db.row_factory = sqlite3.Row
    return db

def init_db():
    db = get_db()
    db.executescript('''
        CREATE TABLE IF NOT EXISTS users (id INTEGER PRIMARY KEY, name TEXT, email TEXT UNIQUE, password TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP);
        CREATE TABLE IF NOT EXISTS analyses (id INTEGER PRIMARY KEY, user_id INTEGER, filename TEXT, original_score INTEGER, improved_score INTEGER, template_used TEXT, matched_keywords TEXT, missing_keywords TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP);
        CREATE TABLE IF NOT EXISTS resumes (id INTEGER PRIMARY KEY, user_id INTEGER, analysis_id INTEGER, file_path TEXT, format TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP);
        CREATE INDEX IF NOT EXISTS idx_user_email ON users(email);
    ''')
    db.commit()
    db.close()
    logger.info("✓ Database initialized")

with app.app_context():
    init_db()

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise Exception("Failed to extract PDF text")

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += '\n' + cell.text
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise Exception("Failed to extract DOCX text")

def extract_contact_info(text):
    emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    phones = re.findall(r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}', text)
    linkedin = re.findall(r'linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
    return {
        'emails': emails[:1] if emails else [], 
        'phones': [p for p in phones if len(re.sub(r'\D', '', p)) >= 10][:1], 
        'linkedin': linkedin[:1] if linkedin else []
    }

def extract_name_from_text(text):
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if lines:
        name = re.sub(r'\b(Mr\.|Mrs\.|Ms\.|Dr\.|Prof\.)\b', '', lines[0], flags=re.IGNORECASE).strip()
        return name if 2 <= len(name.split()) <= 4 else "Professional"
    return "Professional"

def parse_resume_structure(text):
    text_lower = text.lower()
    sections = {
        'experience': bool(re.search(r'\b(experience|work history|employment)\b', text_lower)),
        'education': bool(re.search(r'\b(education|academic|degree|university)\b', text_lower)),
        'skills': bool(re.search(r'\b(skills|technical skills)\b', text_lower)),
        'projects': bool(re.search(r'\b(projects?|portfolio)\b', text_lower)),
        'certifications': bool(re.search(r'\b(certifications?|certificates?)\b', text_lower)),
        'languages': bool(re.search(r'\b(languages?|linguistic)\b', text_lower))
    }
    return {
        'name': extract_name_from_text(text), 
        'contact': extract_contact_info(text), 
        'sections': sections, 
        'full_text': text
    }

def calculate_ats_score(resume_text, job_description):
    resume_lower = resume_text.lower()
    job_lower = job_description.lower()
    stop_words = {'the', 'and', 'for', 'with', 'this', 'that', 'from', 'have', 'will', 'your', 'our'}
    
    job_keywords = {w for w in re.findall(r'\b\w+\b', job_lower) if len(w) > 3 and w not in stop_words}
    resume_keywords = set(re.findall(r'\b\w+\b', resume_lower))
    
    matched = job_keywords.intersection(resume_keywords)
    missing = job_keywords - matched
    match_pct = (len(matched) / len(job_keywords) * 100) if job_keywords else 0
    
    resume_data = parse_resume_structure(resume_text)
    has_contact = len(resume_data['contact']['emails']) > 0
    has_phone = len(resume_data['contact']['phones']) > 0
    sections_found = sum(resume_data['sections'].values())
    structure_score = (has_contact * 10) + (has_phone * 10) + (sections_found * 10)
    
    ats_score = min(100, int((match_pct * 0.65) + (structure_score * 0.35)))
    
    return {
        'score': ats_score, 
        'matched_keywords': sorted(list(matched))[:30], 
        'missing_keywords': sorted(list(missing))[:30], 
        'has_contact': has_contact, 
        'has_phone': has_phone, 
        'sections': resume_data['sections'], 
        'sections_found': sections_found, 
        'contact_info': resume_data['contact'], 
        'keyword_match_rate': round(match_pct, 2), 
        'resume_data': resume_data
    }

def generate_ai_suggestions(resume_text, job_description, score_data):
    suggestions = []
    score = score_data['score']
    
    if score < 50:
        suggestions.append({'type': 'critical', 'title': '🚨 Very Low ATS Score', 'description': f"Score: {score}/100. Major improvements needed."})
    elif score < 70:
        suggestions.append({'type': 'warning', 'title': '⚠️ Low Keyword Match', 'description': f"Matches {score_data['keyword_match_rate']}% of keywords."})
    
    if not score_data['has_contact']:
        suggestions.append({'type': 'critical', 'title': '❌ Missing Email', 'description': 'Add professional email.'})
    
    if not score_data['has_phone']:
        suggestions.append({'type': 'warning', 'title': '📞 Missing Phone', 'description': 'Add contact number.'})
    
    if score >= 80:
        suggestions.append({'type': 'success', 'title': '✅ Excellent Score!', 'description': 'Well-optimized resume.'})
    
    return suggestions

def generate_resume_with_ai(resume_data, job_description, template_style='professional'):
    if not groq_client:
        raise Exception("Groq API not configured. Add GROQ_API_KEY to .env file.")
    
    original_text = resume_data.get('full_text', '')
    original_name = resume_data.get('name', 'Professional')
    original_contact = resume_data.get('contact', {})
    
    contact_email = original_contact.get('emails', ['email@example.com'])[0] if original_contact.get('emails') else 'email@example.com'
    contact_phone = original_contact.get('phones', ['+1234567890'])[0] if original_contact.get('phones') else '+1234567890'
    contact_linkedin = original_contact.get('linkedin', [''])[0] if original_contact.get('linkedin') else ''
    
    prompt = """You are an expert ATS resume writer. IMPROVE this resume using the job description.

**ORIGINAL RESUME:**
""" + original_text[:3500] + """

**JOB DESCRIPTION:**
""" + job_description[:1500] + """

**RULES:**
1. PRESERVE truthful info (names, companies, dates)
2. FIX grammar and spelling
3. ADD keywords from JD naturally
4. ENHANCE bullets with action verbs + metrics
5. Include LinkedIn if available
6. Add Languages section if mentioned

**RETURN THIS JSON (NO MARKDOWN, NO CODE BLOCKS):**
{
    "name": \"""" + original_name + """\",
    "contact": {
        "email": \"""" + contact_email + """\",
        "phone": \"""" + contact_phone + """\",
        "location": "City, State",
        "linkedin": \"""" + contact_linkedin + """\"
    },
    "summary": "2-3 sentences with years of experience, key skills matching JD, quantifiable achievements.",
    "experience": [
        {
            "title": "Actual Job Title",
            "company": "Actual Company",
            "duration": "MM/YYYY - MM/YYYY",
            "achievements": [
                "Led project achieving 30% efficiency gain through implementation of X",
                "Managed team of 5 delivering solutions that increased revenue by $500K",
                "Optimized process reducing costs by 20% using Y technology"
            ]
        }
    ],
    "education": [
        {"degree": "Actual Degree", "institution": "Actual School", "year": "YYYY", "gpa": "X.X (if mentioned)"}
    ],
    "skills": ["Skill1", "Skill2", "Skill3"],
    "certifications": ["Actual Cert 1", "Actual Cert 2"],
    "languages": ["English (Native)", "Spanish (Fluent)"],
    "projects": ["Project name: Description with impact metrics"]
}

QUALITY: Use strong action verbs (Led, Achieved, Implemented). Add metrics (%, $, numbers). Keep bullets 10-20 words. Match JD keywords."""

    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "Expert resume writer. Return ONLY valid JSON. No markdown."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4096
        )
        
        result = response.choices[0].message.content.strip()
        
        # FIXED: Proper markdown cleanup
        if "```json" in result:
            parts = result.split("```json")
            if len(parts) > 1:
                result = parts[1].split("```")[0].strip()
        elif "```" in result:
            parts = result.split("```")
            if len(parts) >= 3:
                result = parts[1].strip()
        
        resume_content = json.loads(result)
        logger.info(f"✓ Resume enhanced: {len(resume_content.get('experience', []))} jobs")
        return resume_content
        
    except json.JSONDecodeError as e:
        logger.error(f"JSON parse error: {e}")
        logger.error(f"Response preview: {result[:500]}")
        raise Exception("AI returned invalid JSON. Please try again.")
    except Exception as e:
        logger.error(f"AI generation error: {e}")
        raise Exception(f"Resume generation failed: {str(e)}")

def create_resume_pdf(resume_content, template_style, filename):
    """FIXED PDF generation - matches preview exactly"""
    try:
        pdf = SimpleDocTemplate(
            filename, 
            pagesize=letter, 
            topMargin=0.5*inch, 
            bottomMargin=0.5*inch, 
            leftMargin=0.7*inch, 
            rightMargin=0.7*inch
        )
        story = []
        styles = getSampleStyleSheet()
        
        # EXACT STYLES TO MATCH PREVIEW
        name_style = ParagraphStyle(
            'Name', 
            parent=styles['Heading1'], 
            fontSize=22, 
            textColor=colors.HexColor('#000000'), 
            alignment=TA_CENTER, 
            spaceAfter=6, 
            fontName='Helvetica-Bold', 
            leading=26
        )
        
        contact_style = ParagraphStyle(
            'Contact', 
            parent=styles['Normal'], 
            fontSize=9, 
            textColor=colors.HexColor('#555555'), 
            alignment=TA_CENTER, 
            spaceAfter=10, 
            fontName='Helvetica'
        )
        
        section_style = ParagraphStyle(
            'Section', 
            parent=styles['Heading2'], 
            fontSize=11, 
            textColor=colors.HexColor('#2d55ff'), 
            spaceAfter=6, 
            spaceBefore=12, 
            fontName='Helvetica-Bold', 
            leading=13
        )
        
        job_title_style = ParagraphStyle(
            'JobTitle', 
            parent=styles['Normal'], 
            fontSize=10, 
            textColor=colors.HexColor('#000000'), 
            fontName='Helvetica-Bold', 
            spaceAfter=2, 
            leading=12
        )
        
        duration_style = ParagraphStyle(
            'Duration', 
            parent=styles['Normal'], 
            fontSize=9, 
            textColor=colors.HexColor('#666666'), 
            fontName='Helvetica', 
            spaceAfter=5, 
            leading=11
        )
        
        bullet_style = ParagraphStyle(
            'Bullet', 
            parent=styles['Normal'], 
            fontSize=9, 
            textColor=colors.HexColor('#333333'), 
            fontName='Helvetica', 
            leftIndent=15, 
            spaceAfter=3, 
            leading=13
        )
        
        body_style = ParagraphStyle(
            'Body', 
            parent=styles['Normal'], 
            fontSize=9, 
            textColor=colors.HexColor('#333333'), 
            fontName='Helvetica', 
            spaceAfter=5, 
            leading=13, 
            alignment=TA_JUSTIFY
        )
        
        # NAME
        story.append(Paragraph(resume_content.get('name', 'Professional').upper(), name_style))
        story.append(Spacer(1, 0.05*inch))
        
        # CONTACT
        contact = resume_content.get('contact', {})
        parts = [
            contact.get('email', ''), 
            contact.get('phone', ''), 
            contact.get('location', '')
        ]
        if contact.get('linkedin'):
            parts.append(contact['linkedin'])
        contact_text = ' | '.join([p for p in parts if p])
        story.append(Paragraph(contact_text, contact_style))
        story.append(Spacer(1, 0.1*inch))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e0e0e0'), spaceAfter=10))
        
        # SUMMARY
        if resume_content.get('summary'):
            story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
            story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#2d55ff'), spaceAfter=6))
            story.append(Paragraph(resume_content['summary'], body_style))
            story.append(Spacer(1, 0.08*inch))
        
        # EXPERIENCE
        if resume_content.get('experience'):
            story.append(Paragraph("PROFESSIONAL EXPERIENCE", section_style))
            story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#2d55ff'), spaceAfter=6))
            
            for i, exp in enumerate(resume_content['experience']):
                story.append(Paragraph(f"{exp.get('title', 'Position')} - {exp.get('company', 'Company')}", job_title_style))
                if exp.get('duration'):
                    story.append(Paragraph(exp['duration'], duration_style))
                for ach in exp.get('achievements', []):
                    story.append(Paragraph(f"• {ach}", bullet_style))
                if i < len(resume_content['experience']) - 1:
                    story.append(Spacer(1, 0.1*inch))
            
            story.append(Spacer(1, 0.08*inch))
        
        # EDUCATION
        if resume_content.get('education'):
            story.append(Paragraph("EDUCATION", section_style))
            story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#2d55ff'), spaceAfter=6))
            
            for edu in resume_content['education']:
                edu_text = f"<b>{edu.get('degree', 'Degree')}</b> - {edu.get('institution', 'School')}"
                if edu.get('year'):
                    edu_text += f" ({edu['year']})"
                if edu.get('gpa'):
                    edu_text += f" | GPA: {edu['gpa']}"
                story.append(Paragraph(edu_text, body_style))
            
            story.append(Spacer(1, 0.08*inch))
        
        # SKILLS
        if resume_content.get('skills') and len(resume_content['skills']) > 0:
            story.append(Paragraph("TECHNICAL SKILLS", section_style))
            story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#2d55ff'), spaceAfter=6))
            story.append(Paragraph(" • ".join(resume_content['skills']), body_style))
            story.append(Spacer(1, 0.08*inch))
        
        # CERTIFICATIONS
        if resume_content.get('certifications') and len(resume_content['certifications']) > 0:
            story.append(Paragraph("CERTIFICATIONS", section_style))
            story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#2d55ff'), spaceAfter=6))
            for cert in resume_content['certifications']:
                story.append(Paragraph(f"• {cert}", bullet_style))
            story.append(Spacer(1, 0.08*inch))
        
        # LANGUAGES
        if resume_content.get('languages') and len(resume_content['languages']) > 0:
            story.append(Paragraph("LANGUAGES", section_style))
            story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#2d55ff'), spaceAfter=6))
            story.append(Paragraph(" • ".join(resume_content['languages']), body_style))
            story.append(Spacer(1, 0.08*inch))
        
        # PROJECTS
        if resume_content.get('projects') and len(resume_content['projects']) > 0:
            story.append(Paragraph("PROJECTS", section_style))
            story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#2d55ff'), spaceAfter=6))
            for proj in resume_content['projects']:
                story.append(Paragraph(f"• {proj}", bullet_style))
        
        pdf.build(story)
        logger.info(f"✓ PDF created: {filename}")
        return True
        
    except Exception as e:
        logger.error(f"PDF error: {e}")
        import traceback
        traceback.print_exc()
        return False

# API ROUTES
FRONTEND_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'frontend')

@app.route('/')
def home():
    try:
        return send_file(os.path.join(FRONTEND_DIR, '1-login.html'))
    except:
        return jsonify({'error': 'Frontend not found'}), 404

@app.route('/<int:step>-<page>.html')
def serve_page(step, page):
    try:
        return send_file(os.path.join(FRONTEND_DIR, f"{step}-{page}.html"))
    except:
        return jsonify({'error': 'Page not found'}), 404

@app.route('/api/register', methods=['POST'])
@limiter.limit("5 per hour")
def register():
    try:
        data = request.json
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        name = data.get('name', '').strip()
        
        if not email or not password or not name or len(password) < 6:
            return jsonify({'error': 'Invalid input'}), 400
        
        db = get_db()
        if db.execute('SELECT id FROM users WHERE email = ?', (email,)).fetchone():
            db.close()
            return jsonify({'error': 'Email exists'}), 400
        
        db.execute('INSERT INTO users (name, email, password) VALUES (?, ?, ?)', 
                   (name, email, generate_password_hash(password)))
        db.commit()
        db.close()
        
        logger.info(f"✓ User registered: {email}")
        return jsonify({'message': 'Success'}), 201
        
    except Exception as e:
        logger.error(f"Registration error: {e}")
        return jsonify({'error': 'Failed'}), 500

@app.route('/api/login', methods=['POST'])
@limiter.limit("10 per minute")
def login():
    try:
        data = request.json
        email = data.get('email', '').strip().lower()
        password = data.get('password', '')
        
        db = get_db()
        user = db.execute('SELECT * FROM users WHERE email = ?', (email,)).fetchone()
        db.close()
        
        if not user or not check_password_hash(user['password'], password):
            return jsonify({'error': 'Invalid credentials'}), 401
        
        session['user_id'] = user['id']
        session['user_email'] = user['email']
        session['user_name'] = user['name']
        session.permanent = True
        
        logger.info(f"✓ User logged in: {email}")
        return jsonify({'message': 'Success', 'user': {'name': user['name'], 'email': user['email']}}), 200
        
    except Exception as e:
        logger.error(f"Login error: {e}")
        return jsonify({'error': 'Failed'}), 500

@app.route('/api/logout', methods=['POST'])
def logout():
    session.clear()
    return jsonify({'message': 'Success'}), 200

@app.route('/api/check-auth', methods=['GET'])
def check_auth():
    if 'user_id' in session:
        try:
            db = get_db()
            user = db.execute('SELECT name, email FROM users WHERE id = ?', (session['user_id'],)).fetchone()
            db.close()
            if user:
                return jsonify({'authenticated': True, 'user': {'name': user['name'], 'email': user['email']}}), 200
        except Exception as e:
            logger.error(f"Auth check error: {e}")
    return jsonify({'authenticated': False}), 200

@app.route('/api/analyze', methods=['POST'])
@limiter.limit("10 per hour")
def analyze_resume():
    filepath = None
    try:
        if 'user_id' not in session:
            return jsonify({'error': 'Auth required'}), 401
        
        if 'resume' not in request.files:
            return jsonify({'error': 'No file'}), 400
        
        file = request.files['resume']
        job_description = request.form.get('job_description', '').strip()
        
        if not file.filename or not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file'}), 400
        
        if len(job_description) < 50:
            return jsonify({'error': 'Job description too short'}), 400
        
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{filename}")
        file.save(filepath)
        
        # Extract text
        if filename.lower().endswith('.pdf'):
            resume_text = extract_text_from_pdf(filepath)
        else:
            resume_text = extract_text_from_docx(filepath)
        
        # Clean up file
        if os.path.exists(filepath):
            os.remove(filepath)
            filepath = None
        
        if len(resume_text) < 100:
            return jsonify({'error': 'Insufficient text'}), 400
        
        # Calculate score
        score_data = calculate_ats_score(resume_text, job_description)
        suggestions = generate_ai_suggestions(resume_text, job_description, score_data)
        
        # Store in session
        session['resume_text'] = resume_text
        session['job_description'] = job_description
        session['resume_data'] = score_data['resume_data']
        session['ats_score'] = score_data['score']
        
        # Save to database
        db = get_db()
        cursor = db.execute(
            'INSERT INTO analyses (user_id, filename, original_score, matched_keywords, missing_keywords) VALUES (?, ?, ?, ?, ?)', 
            (session['user_id'], file.filename, score_data['score'], 
             ','.join(score_data['matched_keywords'][:10]), 
             ','.join(score_data['missing_keywords'][:10]))
        )
        session['analysis_id'] = cursor.lastrowid
        db.commit()
        db.close()
        
        logger.info(f"✓ Analysis complete: {score_data['score']}/100")
        
        return jsonify({
            'score': score_data['score'], 
            'matched_keywords': score_data['matched_keywords'], 
            'missing_keywords': score_data['missing_keywords'], 
            'suggestions': suggestions, 
            'analysis': {
                'has_contact': score_data['has_contact'], 
                'has_phone': score_data['has_phone'], 
                'sections_found': score_data['sections_found'], 
                'keyword_match_rate': score_data['keyword_match_rate']
            }
        }), 200
        
    except Exception as e:
        if filepath and os.path.exists(filepath):
            os.remove(filepath)
        logger.error(f"Analysis error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate-resume', methods=['POST'])
@limiter.limit("5 per hour")
def generate_resume():
    try:
        if 'user_id' not in session:
            return jsonify({'error': 'Auth required'}), 401
        
        template_style = request.json.get('template_style', 'professional')
        resume_data = session.get('resume_data')
        job_description = session.get('job_description')
        
        if not resume_data or not job_description:
            return jsonify({'error': 'No data found'}), 400
        
        logger.info(f"Generating resume with {template_style} template...")
        resume_content = generate_resume_with_ai(resume_data, job_description, template_style)
        
        session['generated_resume'] = resume_content
        session['resume_template'] = template_style
        
        logger.info("✓ Resume generated successfully")
        return jsonify({'message': 'Success', 'resume_content': resume_content}), 200
        
    except Exception as e:
        logger.error(f"Generation error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/download-resume', methods=['POST'])
def download_resume():
    try:
        if 'user_id' not in session:
            return jsonify({'error': 'Auth required'}), 401
        
        resume_content = session.get('generated_resume')
        if not resume_content:
            return jsonify({'error': 'No resume to download'}), 400
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{app.config['GENERATED_FOLDER']}/resume_{session['user_id']}_{timestamp}.pdf"
        
        if create_resume_pdf(resume_content, session.get('resume_template', 'professional'), filename):
            db = get_db()
            db.execute(
                'INSERT INTO resumes (user_id, analysis_id, file_path, format) VALUES (?, ?, ?, ?)', 
                (session['user_id'], session.get('analysis_id'), filename, 'pdf')
            )
            db.commit()
            db.close()
            
            logger.info(f"✓ PDF download ready: {filename}")
            return send_file(filename, as_attachment=True, download_name='ATS_Optimized_Resume.pdf')
        else:
            return jsonify({'error': 'PDF creation failed'}), 500
            
    except Exception as e:
        logger.error(f"Download error: {e}")
        return jsonify({'error': 'Failed'}), 500

@app.route('/api/recalculate-score', methods=['POST'])
def recalculate_score():
    try:
        if 'user_id' not in session:
            return jsonify({'error': 'Auth required'}), 401
        
        resume_content = session.get('generated_resume')
        job_description = session.get('job_description')
        
        if not resume_content or not job_description:
            return jsonify({'error': 'No data'}), 400
        
        # Reconstruct resume text from JSON
        resume_text = f"{resume_content['name']}\n"
        resume_text += f"{json.dumps(resume_content['contact'])}\n"
        resume_text += f"{resume_content.get('summary', '')}\n"
        
        for exp in resume_content.get('experience', []):
            resume_text += f"{exp.get('title', '')} {exp.get('company', '')}\n"
            resume_text += f"{exp.get('duration', '')}\n"
            resume_text += '\n'.join(exp.get('achievements', [])) + "\n"
        
        for edu in resume_content.get('education', []):
            resume_text += f"{edu.get('degree', '')} {edu.get('institution', '')}\n"
        
        resume_text += ' '.join(resume_content.get('skills', []))
        
        # Calculate new score
        score_data = calculate_ats_score(resume_text, job_description)
        
        # Update database
        if session.get('analysis_id'):
            db = get_db()
            db.execute(
                'UPDATE analyses SET improved_score = ?, template_used = ? WHERE id = ?', 
                (score_data['score'], session.get('resume_template'), session['analysis_id'])
            )
            db.commit()
            db.close()
        
        improvement = score_data['score'] - session.get('ats_score', 0)
        logger.info(f"✓ Score recalculated: {score_data['score']}/100 (improvement: +{improvement})")
        
        return jsonify({
            'new_score': score_data['score'], 
            'improvement': improvement
        }), 200
        
    except Exception as e:
        logger.error(f"Recalc error: {e}")
        return jsonify({'error': 'Failed'}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'groq_configured': groq_client is not None,
        'timestamp': datetime.now().isoformat()
    }), 200

# ERROR HANDLERS
@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Not found'}), 404

@app.errorhandler(413)
def file_too_large(e):
    return jsonify({'error': 'File too large (max 5MB)'}), 413

@app.errorhandler(429)
def rate_limit_exceeded(e):
    return jsonify({'error': 'Rate limit exceeded'}), 429

@app.errorhandler(500)
def internal_error(e):
    logger.error(f"Internal error: {e}")
    return jsonify({'error': 'Server error'}), 500

if __name__ == '__main__':
    print("\n" + "="*70)
    print("🚀 Smart ATS Resume Checker v4.0 - ALL ERRORS FIXED")
    print("="*70)
    print(f"📍 Server: http://localhost:5000")
    print(f"🤖 AI Engine: Groq (Llama 3.3 70B)")
    print(f"💾 Database: database/ats_checker.db")
    print("="*70)
    print("✅ ENHANCEMENTS:")
    print("   ✓ Fixed PDF formatting (matches preview exactly)")
    print("   ✓ Added LinkedIn, Languages, Projects sections")
    print("   ✓ Enhanced grammar correction via AI")
    print("   ✓ Better JD keyword integration")
    print("   ✓ Improved ATS scoring algorithm")
    print("   ✓ Fixed all JSON parsing errors")
    print("   ✓ Fixed triple-quoted string syntax errors")
    print("   ✓ Better error handling & logging")
    print("   ✓ Proper file cleanup")
    print("="*70)
    
    # Startup checks
    issues = []
    if not groq_client:
        issues.append("⚠️  GROQ_API_KEY not configured - AI features disabled")
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        issues.append(f"⚠️  Creating upload folder: {app.config['UPLOAD_FOLDER']}")
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    if not os.path.exists(app.config['GENERATED_FOLDER']):
        issues.append(f"⚠️  Creating generated folder: {app.config['GENERATED_FOLDER']}")
        os.makedirs(app.config['GENERATED_FOLDER'], exist_ok=True)
    
    if issues:
        print("\n" + "="*70)
        print("STARTUP WARNINGS:")
        for issue in issues:
            print(f"  {issue}")
        print("="*70)
    
    print("\n🎯 Server ready! All syntax errors fixed.\n")
    app.run(debug=True, host='0.0.0.0', port=5000, threaded=True)